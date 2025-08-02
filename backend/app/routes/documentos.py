from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from docx import Document
from datetime import datetime
import os
import unicodedata
import re
import subprocess

from app.database import get_db
from app.models.advogado import Advogado
from app.google_drive_uploader import upload_pdf_to_drive
from app.zapsign import enviar_para_assinatura  # ✅ ZapSign

router = APIRouter()

# ✅ Normaliza nomes (slug)
def slugify(texto: str) -> str:
    texto = unicodedata.normalize('NFD', texto)
    texto = texto.encode('ascii', 'ignore').decode('utf-8')
    texto = re.sub(r'\s+', '_', texto)
    texto = re.sub(r'[^a-zA-Z0-9_]', '', texto)
    return texto.lower()

# ✅ Converte DOCX para PDF
def converter_docx_para_pdf(caminho_docx: str, pasta_saida: str):
    try:
        subprocess.run([
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "--headless", "--convert-to", "pdf",
            "--outdir", pasta_saida, caminho_docx
        ], check=True)
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Erro na conversão com LibreOffice: {e}")

# ✅ Preenche modelo com dados (corrigido para evitar perda de campos como 'Brasileira trans')
def preencher_documento(modelo_path: str, dados: dict, nome_saida: str) -> str:
    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Modelo não encontrado: {modelo_path}")

    doc = Document(modelo_path)

    for par in doc.paragraphs:
        texto = par.text
        for chave, valor in dados.items():
            valor_str = str(valor) if valor is not None else ""
            texto = texto.replace(f"{{{{ {chave} }}}}", valor_str)
            texto = texto.replace(f"{{{{{chave}}}}}", valor_str)
        if texto != par.text:
            par.clear()
            par.add_run(texto)

    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    nome_cliente_slug = slugify(dados['nome'])
    nome_base = f"{nome_saida}_{nome_cliente_slug}_{timestamp}"

    pasta_saida = os.path.join(os.getcwd(), "documentos_gerados")
    os.makedirs(pasta_saida, exist_ok=True)

    caminho_docx = os.path.join(pasta_saida, f"{nome_base}.docx")
    caminho_pdf = os.path.join(pasta_saida, f"{nome_base}.pdf")

    doc.save(caminho_docx)
    converter_docx_para_pdf(caminho_docx, pasta_saida)

    if not os.path.exists(caminho_pdf):
        raise RuntimeError(f"❌ PDF não foi gerado: {caminho_pdf}")

    return f"{nome_base}.pdf"

# ✅ Rota principal
@router.post("/gerar-documentos")
def gerar_documentos(dados: dict, db: Session = Depends(get_db)):
    try:
        print("📨 Iniciando geração de documentos...")

        base = os.path.join(os.getcwd(), "modelos", "vitor")
        contrato_path = os.path.join(base, "modelo_contrato_vitor.docx")
        procuracao_path = os.path.join(base, "modelo_procuracao_vitor.docx")

        contrato = preencher_documento(contrato_path, dados, "contrato")
        procuracao = preencher_documento(procuracao_path, dados, "procuracao")

        caminho_procuracao_pdf = os.path.join(os.getcwd(), "documentos_gerados", procuracao)
        url_documento = upload_pdf_to_drive(caminho_procuracao_pdf, procuracao)

        # 🔍 Busca o advogado
        advogado = db.query(Advogado).filter(Advogado.usuario == dados.get("usuario_advogado")).first()
        if not advogado or not advogado.api_key_zapsign:
            raise HTTPException(status_code=400, detail="Advogado não encontrado ou sem chave ZapSign")

        print("📤 Enviando para ZapSign...")
        print("🔗 Documento:", url_documento)
        print("📱 Telefone:", dados.get("telefone"))
        print("🔐 API Key:", advogado.api_key_zapsign[:5] + "..." + advogado.api_key_zapsign[-4:])

        # 📤 Envia para ZapSign
        resposta_zapsign = enviar_para_assinatura(
            nome_cliente=dados["nome"],
            telefone_cliente=dados["telefone"],
            url_documento=url_documento,
            nome_documento="Procuração PJMOL",
            api_key=advogado.api_key_zapsign
        )

        print("📥 Resposta ZapSign:", resposta_zapsign)

        return {
            "mensagem": "Documentos gerados com sucesso!",
            "contrato_pdf": contrato,
            "procuracao_pdf": procuracao,
            "url_procuracao_drive": url_documento,
            "zapsign": resposta_zapsign
        }

    except Exception as e:
        print("❌ Erro ao gerar ou enviar documentos:", e)
        raise HTTPException(status_code=500, detail=str(e))

# ✅ Servir documento
@router.get("/documentos/{nome_arquivo}")
def servir_documento(nome_arquivo: str):
    caminho = os.path.join(os.getcwd(), "documentos_gerados", nome_arquivo)
    if not os.path.exists(caminho):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")

    if nome_arquivo.endswith(".pdf"):
        media_type = "application/pdf"
    elif nome_arquivo.endswith(".docx"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        media_type = "application/octet-stream"

    return FileResponse(caminho, media_type=media_type)
