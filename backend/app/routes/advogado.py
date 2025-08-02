from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from docx import Document
from datetime import datetime
import os
import unicodedata
import re
import subprocess

from pydantic import BaseModel
from app.database import get_db
from app.models.advogado import Advogado
from app.utils.security import gerar_hash_senha  # ✅ segurança de senha

router = APIRouter()

# ✅ Normaliza strings (sem acento ou espaço)
def slugify(texto: str) -> str:
    texto = unicodedata.normalize('NFD', texto)
    texto = texto.encode('ascii', 'ignore').decode('utf-8')
    texto = re.sub(r'\s+', '_', texto)
    texto = re.sub(r'[^a-zA-Z0-9_]', '', texto)
    return texto.lower()

# ✅ Converte .docx → .pdf usando LibreOffice
def converter_docx_para_pdf(caminho_docx: str, pasta_saida: str):
    try:
        subprocess.run([
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", pasta_saida,
            caminho_docx
        ], check=True)
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Erro na conversão com LibreOffice: {e}")

# ✅ Substitui placeholders no .docx e gera PDF (corrigido)
def preencher_documento(modelo_path: str, dados: dict, nome_saida: str) -> str:
    if not os.path.exists(modelo_path):
        raise FileNotFoundError(f"Modelo não encontrado: {modelo_path}")

    doc = Document(modelo_path)

    for par in doc.paragraphs:
        texto = par.text
        for chave, valor in dados.items():
            valor_str = str(valor) if valor is not None else ""
            texto = texto.replace(f"{{{{ {chave} }}}}", valor_str)
            texto = texto.replace(f"{{{{{chave}}}}}", valor_str)
        if texto != par.text:
            par.clear()
            par.add_run(texto)

    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    nome_cliente_slug = slugify(dados['nome'])
    nome_base = f"{nome_saida}_{nome_cliente_slug}_{timestamp}"

    pasta_saida = os.path.join(os.getcwd(), "documentos_gerados")
    os.makedirs(pasta_saida, exist_ok=True)

    caminho_docx = os.path.join(pasta_saida, f"{nome_base}.docx")
    caminho_pdf = os.path.join(pasta_saida, f"{nome_base}.pdf")

    doc.save(caminho_docx)
    converter_docx_para_pdf(caminho_docx, pasta_saida)

    if not os.path.exists(caminho_pdf):
        raise RuntimeError(f"❌ PDF não foi gerado: {caminho_pdf}")

    return f"{nome_base}.pdf"

# ✅ Rota principal para geração de documentos
@router.post("/gerar-documentos")
def gerar_documentos(dados: dict, db: Session = Depends(get_db)):
    try:
        usuario = dados.get("usuario_advogado")
        if not usuario:
            raise HTTPException(status_code=400, detail="Usuário do advogado não informado")

        advogado = db.query(Advogado).filter(Advogado.usuario == usuario).first()
        if not advogado:
            raise HTTPException(status_code=404, detail="Advogado não encontrado")

        base = os.path.join(os.getcwd(), "modelos", "vitor")
        contrato_path = os.path.join(base, "modelo_contrato_vitor.docx")
        procuracao_path = os.path.join(base, "modelo_procuracao_vitor.docx")

        contrato = preencher_documento(contrato_path, dados, "contrato")
        procuracao = preencher_documento(procuracao_path, dados, "procuracao")

        return {
            "mensagem": "Documentos gerados com sucesso!",
            "contrato_pdf": contrato,
            "procuracao_pdf": procuracao
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ✅ Rota para baixar documento
@router.get("/documentos/{nome_arquivo}")
def servir_documento(nome_arquivo: str):
    caminho = os.path.join(os.getcwd(), "documentos_gerados", nome_arquivo)
    if not os.path.exists(caminho):
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")

    if nome_arquivo.endswith(".pdf"):
        media_type = "application/pdf"
    elif nome_arquivo.endswith(".docx"):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        media_type = "application/octet-stream"

    return FileResponse(caminho, media_type=media_type)

# ✅ Rota para buscar advogado pelo nome de usuário
@router.get("/advogados/usuario/{usuario}")
def consultar_advogado_por_usuario(usuario: str, db: Session = Depends(get_db)):
    advogado = db.query(Advogado).filter(Advogado.usuario == usuario).first()
    if not advogado:
        raise HTTPException(status_code=404, detail="Advogado não encontrado")

    return {
        "id": advogado.id,
        "nome_completo": advogado.nome_completo,
        "usuario": advogado.usuario,
        "oab": advogado.oab,
        "email": advogado.email,
        "telefone": advogado.telefone,
        "api_key_zapsign": advogado.api_key_zapsign  # ✅ adicionado aqui
    }

# ✅ Modelo de entrada para cadastro de advogado
class AdvogadoCreate(BaseModel):
    nome_completo: str
    oab: str
    email: str
    telefone: str
    usuario: str
    senha: str

# ✅ Rota para cadastrar advogado
@router.post("/advogados/")
def criar_advogado(dados: AdvogadoCreate, db: Session = Depends(get_db)):
    if db.query(Advogado).filter(Advogado.usuario == dados.usuario).first():
        raise HTTPException(status_code=400, detail="Usuário já existe")

    novo_advogado = Advogado(
        nome_completo=dados.nome_completo,
        oab=dados.oab,
        email=dados.email,
        telefone=dados.telefone,
        usuario=dados.usuario,
        senha_hash=gerar_hash_senha(dados.senha)
    )

    db.add(novo_advogado)
    db.commit()
    db.refresh(novo_advogado)

    return {
        "mensagem": "Advogado cadastrado com sucesso!",
        "id": novo_advogado.id
    }

# ✅ Rota para listar todos os advogados
@router.get("/advogados/")
def listar_advogados(db: Session = Depends(get_db)):
    advogados = db.query(Advogado).all()
    return [
        {
            "id": adv.id,
            "nome_completo": adv.nome_completo,
            "usuario": adv.usuario,
            "email": adv.email,
            "telefone": adv.telefone,
            "oab": adv.oab,
            "api_key_zapsign": adv.api_key_zapsign  # também incluído aqui se desejar
        }
        for adv in advogados
    ]

# ✅ Rota para atualizar a chave da ZapSign
class AtualizarChaveZapSign(BaseModel):
    api_key_zapsign: str

@router.put("/advogados/usuario/{usuario}")
def atualizar_api_key(usuario: str, dados: AtualizarChaveZapSign, db: Session = Depends(get_db)):
    advogado = db.query(Advogado).filter(Advogado.usuario == usuario).first()
    if not advogado:
        raise HTTPException(status_code=404, detail="Advogado não encontrado")

    advogado.api_key_zapsign = dados.api_key_zapsign
    db.commit()
    db.refresh(advogado)

    return {
        "mensagem": "Chave ZapSign atualizada com sucesso",
        "usuario": advogado.usuario,
        "nova_api_key": advogado.api_key_zapsign
    }
